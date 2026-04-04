from __future__ import annotations

from pathlib import Path


def _must_have_docx() -> None:
    try:
        import docx  # noqa: F401
    except Exception as exc:
        raise SystemExit("python-docx no instalado. Ejecuta: pip install python-docx>=1.1.2") from exc


def generate_carta_template(out_dir: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_heading("NOMBRE DE LA FIRMA", level=0)
    doc.add_paragraph("LOGO: {{firm_logo}}")
    doc.add_paragraph("CARTA DE CONTROL INTERNO (BORRADOR)")
    doc.add_paragraph("Destinatario: {{recipient}}")
    doc.add_paragraph("Asunto: {{subject}}")
    doc.add_paragraph("Entidad: {{company_name}}")
    doc.add_paragraph("Periodo: {{period_end}}")
    doc.add_paragraph("Leyenda: {{borrador_badge}}")
    doc.add_heading("Hallazgos", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Hallazgo"
    table.rows[0].cells[1].text = "Riesgo"
    table.rows[0].cells[2].text = "Recomendacion"
    table.rows[0].cells[3].text = "Respuesta gerencia"
    doc.add_paragraph("Firma socio: {{signature_partner}}")

    out = out_dir / "carta_control_interno.docx"
    doc.save(out)
    return out


def generate_niif_template(out_dir: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_heading("NOMBRE DE LA FIRMA", level=0)
    doc.add_paragraph("LOGO: {{firm_logo}}")
    doc.add_heading("BORRADOR DE NOTAS A LOS EEFF - NIIF PARA PYMES", level=1)
    doc.add_paragraph("Entidad: {{company_name}}")
    doc.add_paragraph("Periodo: {{period_end}}")
    doc.add_paragraph("Version NIIF para PYMES: {{ifrs_for_smes_version}}")
    doc.add_paragraph("Adopcion anticipada: {{early_adoption}}")
    doc.add_paragraph("Estado: {{borrador_badge}}")
    doc.add_page_break()
    doc.add_heading("Indice", level=1)
    doc.add_paragraph("- Informacion general")
    doc.add_paragraph("- Base de preparacion")
    doc.add_paragraph("- Politicas contables")
    doc.add_paragraph("- Juicios y estimaciones")
    doc.add_paragraph("- Notas especificas")
    doc.add_paragraph("- Cierre")
    doc.add_page_break()
    doc.add_heading("Bloque de aprobacion", level=2)
    doc.add_paragraph("Aprobacion: {{approval_block}}")
    doc.add_paragraph("Representante legal: {{signature_legal_rep}}")
    doc.add_paragraph("Contador: {{signature_accountant}}")

    out = out_dir / "notas_niif_pymes.docx"
    doc.save(out)
    return out


def main() -> None:
    _must_have_docx()
    out_dir = Path(__file__).resolve().parents[1] / "templates" / "docx"
    out_dir.mkdir(parents=True, exist_ok=True)
    carta = generate_carta_template(out_dir)
    niif = generate_niif_template(out_dir)
    print(f"Template generado: {carta}")
    print(f"Template generado: {niif}")


if __name__ == "__main__":
    main()
