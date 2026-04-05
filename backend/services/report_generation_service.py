from __future__ import annotations

import hashlib
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any
import yaml

from backend.repositories.file_repository import list_area_codes, read_area_yaml, read_hallazgos, read_perfil
from backend.services.rag_chat_service import generate_judgement_response

ROOT = Path(__file__).resolve().parents[2]
CLIENTES_ROOT = ROOT / "data" / "clientes"
DOCX_TEMPLATES_DIR = ROOT / "backend" / "templates" / "docx"
TEMPLATE_CONTRACTS_PATH = DOCX_TEMPLATES_DIR / "template_contracts.yaml"
TEMPLATE_PROFILES_PATH = DOCX_TEMPLATES_DIR / "template_profiles.yaml"


def _safe_text(value: Any) -> str:
    return str(value or "").strip()


def _report_dir(cliente_id: str) -> Path:
    out = CLIENTES_ROOT / cliente_id / "reportes"
    out.mkdir(parents=True, exist_ok=True)
    return out


def _normalize_recipient(recipient: str) -> str:
    text = _safe_text(recipient)
    return text or "Gerencia General"


def _default_contract(document_type: str) -> dict[str, Any]:
    return {
        "template_version": "v1-default",
        "document_type": document_type,
        "placeholders_supported": [],
        "required_sections": [],
        "optional_sections": [],
    }


def _deep_merge(base: dict[str, Any], override: dict[str, Any]) -> dict[str, Any]:
    out = dict(base)
    for key, value in override.items():
        if isinstance(value, dict) and isinstance(out.get(key), dict):
            out[key] = _deep_merge(out.get(key) or {}, value)
        else:
            out[key] = value
    return out


def _load_template_profile(cliente_id: str, document_type: str) -> dict[str, Any]:
    base: dict[str, Any] = {"document_type": document_type}
    if TEMPLATE_PROFILES_PATH.exists():
        try:
            profiles = yaml.safe_load(TEMPLATE_PROFILES_PATH.read_text(encoding="utf-8")) or {}
        except Exception:
            profiles = {}
        if isinstance(profiles, dict):
            base = _deep_merge(base, profiles.get("base", {}) if isinstance(profiles.get("base"), dict) else {})
            firma = profiles.get("firmas", {}) if isinstance(profiles.get("firmas"), dict) else {}
            # v1: firma por defecto; en v2 se puede enrutar por org/tenant.
            base = _deep_merge(base, firma.get("default", {}) if isinstance(firma.get("default"), dict) else {})

    cliente_override_path = CLIENTES_ROOT / cliente_id / "template_overrides.yaml"
    if cliente_override_path.exists():
        try:
            raw = yaml.safe_load(cliente_override_path.read_text(encoding="utf-8")) or {}
        except Exception:
            raw = {}
        if isinstance(raw, dict):
            base = _deep_merge(base, raw.get(document_type, {}) if isinstance(raw.get(document_type), dict) else {})
    return base


def _load_template_contract(document_type: str, *, cliente_id: str | None = None) -> dict[str, Any]:
    if not TEMPLATE_CONTRACTS_PATH.exists():
        return _default_contract(document_type)
    try:
        data = yaml.safe_load(TEMPLATE_CONTRACTS_PATH.read_text(encoding="utf-8")) or {}
    except Exception:
        return _default_contract(document_type)
    if not isinstance(data, dict):
        return _default_contract(document_type)
    raw = data.get(document_type) or {}
    if not isinstance(raw, dict):
        return _default_contract(document_type)
    contract = _default_contract(document_type)
    contract["template_version"] = _safe_text(raw.get("template_version")) or contract["template_version"]
    contract["document_type"] = _safe_text(raw.get("document_type")) or document_type
    for key in ["placeholders_supported", "required_sections", "optional_sections"]:
        value = raw.get(key)
        if isinstance(value, list):
            contract[key] = [str(x).strip() for x in value if str(x).strip()]
    if cliente_id:
        profile = _load_template_profile(cliente_id, document_type)
        contract = _deep_merge(contract, profile)
    return contract


def _generation_meta_fallback(
    *,
    generated_at: datetime,
    requested_by: str,
    input_payload: dict[str, Any],
    document_type: str,
    template_mode: str,
    template_version: str,
    contract: dict[str, Any],
) -> dict[str, Any]:
    return {
        "source": "fallback",
        "provider": "fallback",
        "model": "",
        "prompt_id": "",
        "prompt_version": "",
        "document_type": document_type,
        "template_mode": template_mode,
        "template_version": template_version,
        "placeholders_supported": contract.get("placeholders_supported") or [],
        "required_sections": contract.get("required_sections") or [],
        "optional_sections": contract.get("optional_sections") or [],
        "generated_at": generated_at,
        "requested_by": requested_by,
        "input_payload": input_payload,
    }


def _generation_meta_llm(
    *,
    rag: dict[str, Any],
    generated_at: datetime,
    requested_by: str,
    input_payload: dict[str, Any],
    document_type: str,
    template_mode: str,
    template_version: str,
    contract: dict[str, Any],
) -> dict[str, Any]:
    prompt_meta = rag.get("prompt_meta") or {}
    return {
        "source": "llm",
        "provider": _safe_text(rag.get("provider")) or "llm",
        "model": _safe_text(rag.get("model")),
        "prompt_id": _safe_text(prompt_meta.get("prompt_id")),
        "prompt_version": _safe_text(prompt_meta.get("prompt_version")),
        "document_type": document_type,
        "template_mode": template_mode,
        "template_version": template_version,
        "placeholders_supported": contract.get("placeholders_supported") or [],
        "required_sections": contract.get("required_sections") or [],
        "optional_sections": contract.get("optional_sections") or [],
        "generated_at": generated_at,
        "requested_by": requested_by,
        "input_payload": input_payload,
    }


def _build_artifact(*, artifact_type: str, artifact_path: Path, payload: bytes, template_version: str) -> dict[str, Any]:
    return {
        "artifact_type": artifact_type,
        "artifact_path": str(artifact_path),
        "artifact_hash": hashlib.sha256(payload).hexdigest(),
        "template_version": template_version,
        "size_bytes": len(payload),
    }


def _load_docx_document(template_filename: str, *, contract: dict[str, Any]) -> tuple[Any, str, str]:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx no esta disponible. Instala dependencia para renderer DOCX.") from exc

    template_path = DOCX_TEMPLATES_DIR / template_filename
    if template_path.exists():
        version = _safe_text(contract.get("template_version")) or f"v1-template:{template_filename}"
        return Document(str(template_path)), version, "custom"
    version = _safe_text(contract.get("template_version")) or "v1-default"
    return Document(), version, "default"


def collect_findings(cliente_id: str, *, max_findings: int) -> list[dict[str, str]]:
    findings: list[dict[str, str]] = []
    for area_code in list_area_codes(cliente_id):
        area = read_area_yaml(cliente_id, area_code)
        area_name = _safe_text(area.get("nombre")) or f"Area {area_code}"
        hallazgos = area.get("hallazgos_abiertos")
        if not isinstance(hallazgos, list):
            continue
        for item in hallazgos:
            if not isinstance(item, dict):
                continue
            findings.append(
                {
                    "id": _safe_text(item.get("id")) or f"{area_code}-OBS",
                    "titulo": _safe_text(item.get("titulo")) or f"Hallazgo {area_name}",
                    "area": area_name,
                    "prioridad": _safe_text(item.get("prioridad")) or "media",
                    "condicion": _safe_text(item.get("descripcion")) or "[[PENDIENTE]] condicion no documentada",
                    "criterio": _safe_text(item.get("aseveracion_afectada"))
                    or "[[PENDIENTE]] criterio especifico de control interno",
                    "causa": _safe_text(item.get("causa")) or "[[PENDIENTE]] causa raiz por documentar",
                    "efecto": _safe_text(item.get("efecto")) or "[[PENDIENTE]] efecto o riesgo no cuantificado",
                    "recomendacion": _safe_text(item.get("recomendacion"))
                    or "[[PENDIENTE]] recomendacion especifica con plan de accion",
                    "respuesta_gerencia": _safe_text(item.get("respuesta_administracion")),
                }
            )
            if len(findings) >= max_findings:
                return findings

    raw_hallazgos = read_hallazgos(cliente_id).strip()
    if raw_hallazgos and len(findings) < max_findings:
        blocks = [b.strip() for b in raw_hallazgos.split("## ") if b.strip()]
        for idx, block in enumerate(blocks, start=1):
            lines = [x.strip() for x in block.splitlines() if x.strip()]
            title = lines[0] if lines else f"Hallazgo {idx}"
            body = " ".join(lines[1:]) if len(lines) > 1 else title
            findings.append(
                {
                    "id": f"HAL-{idx:02d}",
                    "titulo": title,
                    "area": "General",
                    "prioridad": "media",
                    "condicion": body[:600] or "[[PENDIENTE]] condicion no documentada",
                    "criterio": "[[PENDIENTE]] criterio normativo o politico aplicable",
                    "causa": "[[PENDIENTE]] causa raiz",
                    "efecto": "[[PENDIENTE]] efecto o riesgo",
                    "recomendacion": "[[PENDIENTE]] recomendacion y responsable",
                    "respuesta_gerencia": "",
                }
            )
            if len(findings) >= max_findings:
                break
    return findings


def _normalize_internal_control_findings(findings: list[dict[str, str]]) -> list[dict[str, Any]]:
    normalized: list[dict[str, Any]] = []
    for idx, raw in enumerate(findings, start=1):
        if not isinstance(raw, dict):
            continue
        prioridad = _safe_text(raw.get("prioridad")).lower() or "media"
        categoria = _safe_text(raw.get("categoria")).lower()
        impacto_material = str(raw.get("impacto_material", "")).strip().lower() in {"true", "1", "si", "yes"}
        if categoria not in {"mayor_interes", "control_interno"}:
            categoria = "mayor_interes" if prioridad == "alta" or impacto_material else "control_interno"

        observacion = _safe_text(raw.get("observacion")) or _safe_text(raw.get("condicion")) or "[[PENDIENTE]] observacion no documentada"
        recomendacion = _safe_text(raw.get("recomendacion")) or "[[PENDIENTE]] recomendacion no documentada"
        comentarios = _safe_text(raw.get("comentarios_administracion")) or _safe_text(raw.get("respuesta_gerencia")) or "[[PENDIENTE]]"
        antecedentes = _safe_text(raw.get("antecedentes"))
        normalized.append(
            {
                "id": _safe_text(raw.get("id")) or f"HAL-{idx:02d}",
                "titulo": _safe_text(raw.get("titulo")) or f"Hallazgo {idx}",
                "categoria": categoria,
                "antecedentes": antecedentes,
                "observacion": observacion,
                "recomendacion": recomendacion,
                "comentarios_administracion": comentarios,
                "prioridad": prioridad,
                "fuentes": raw.get("fuentes") if isinstance(raw.get("fuentes"), list) else [],
                "area": _safe_text(raw.get("area")),
            }
        )
    return normalized


def _split_findings_by_category(findings: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    major: list[dict[str, Any]] = []
    control: list[dict[str, Any]] = []
    for item in findings:
        if not isinstance(item, dict):
            continue
        if _safe_text(item.get("categoria")) == "mayor_interes":
            major.append(item)
        else:
            control.append(item)
    return major, control


def build_internal_control_letter(
    *,
    company_name: str,
    period_end: str,
    recipient: str,
    findings: list[dict[str, str]],
    include_management_response: bool,
) -> dict[str, Any]:
    normalized_findings = _normalize_internal_control_findings(findings)
    major_findings, control_findings = _split_findings_by_category(normalized_findings)
    today = datetime.now(timezone.utc)
    full_company = _safe_text(company_name) or "[[PENDIENTE]] entidad"
    date_label = today.strftime("%d de %B del %Y")
    recipient_line = _safe_text(recipient) or "Gerencia General"
    return {
        "header": {
            "title": "Informe de Control Interno",
            "date": datetime.now(timezone.utc).date().isoformat(),
            "date_label": date_label,
            "recipient": recipient,
            "company_name": company_name,
            "period_end": period_end,
        },
        "cover": {
            "company_name": full_company,
            "report_title": "Informe de Control Interno",
            "period_end": period_end,
            "ruc": "[[PENDIENTE]]",
        },
        "contenido": {
            "major_interest_titles": [f.get("titulo") for f in major_findings],
            "control_titles": [f.get("titulo") for f in control_findings],
        },
        "letter_intro": {
            "city_and_date": f"Quito, {date_label}",
            "recipient_block": f"A {recipient_line},\n{full_company}",
            "saludo": "Estimados Señores:",
            "paragraphs": [
                (
                    f"En la planeación y ejecución de nuestra auditoría de los estados financieros de {full_company}, "
                    f"al {period_end}, consideramos su estructura de control interno para definir la naturaleza, oportunidad "
                    "y extensión de nuestros procedimientos."
                ),
                (
                    "La Administración de la Compañía es responsable por el establecimiento de los controles internos que "
                    "considere necesarios para la preparación de estados financieros libres de errores significativos."
                ),
                (
                    "Nuestra revisión de control interno tuvo un propósito limitado al diseño de procedimientos de auditoría, "
                    "por lo que no constituye una opinión sobre la efectividad integral del sistema de control interno."
                ),
                (
                    "Debido a las limitaciones inherentes de cualquier sistema de control interno, pueden existir deficiencias "
                    "que no hayan sido identificadas en el alcance de nuestro trabajo."
                ),
            ],
        },
        "intro": (
            "En la planeación y ejecución de nuestra auditoría consideramos la estructura de control interno con un propósito "
            "limitado al diseño de procedimientos."
        ),
        "responsibility": "La Administración es responsable del establecimiento y mantenimiento del control interno.",
        "limitations": "La revisión no constituye una opinión sobre efectividad integral del control interno.",
        "major_interest_findings": major_findings,
        "internal_control_findings": control_findings,
        "findings": normalized_findings,
        "include_management_response": include_management_response,
        "closing": {
            "body": "Agradecemos las facilidades brindadas por la Administración y quedamos atentos para apoyar el seguimiento de los planes de acción.",
            "signature_name": "[[PENDIENTE]] Nombre del Socio",
            "signature_role": "Socio de Auditoría",
            "firm_name": "[[PENDIENTE]] Firma Auditora",
            "copy_to": "[[PENDIENTE]] Contador",
            "draft_legend": "BORRADOR – Documento sujeto a revisión interna",
        },
    }


def render_internal_control_letter_markdown(document: dict[str, Any]) -> str:
    header = document.get("header", {}) if isinstance(document.get("header"), dict) else {}
    cover = document.get("cover", {}) if isinstance(document.get("cover"), dict) else {}
    contenido = document.get("contenido", {}) if isinstance(document.get("contenido"), dict) else {}
    letter_intro = document.get("letter_intro", {}) if isinstance(document.get("letter_intro"), dict) else {}
    major_findings = (
        document.get("major_interest_findings") if isinstance(document.get("major_interest_findings"), list) else []
    )
    control_findings = (
        document.get("internal_control_findings") if isinstance(document.get("internal_control_findings"), list) else []
    )
    findings = document.get("findings", []) if isinstance(document.get("findings"), list) else major_findings + control_findings
    closing = document.get("closing", {}) if isinstance(document.get("closing"), dict) else {}
    include_response = bool(document.get("include_management_response"))
    lines: list[str] = []
    lines.append(_safe_text(cover.get("company_name")) or _safe_text(header.get("company_name")))
    lines.append(_safe_text(cover.get("report_title")) or _safe_text(header.get("title")) or "Informe de Control Interno")
    lines.append(f"Al {_safe_text(cover.get('period_end')) or _safe_text(header.get('period_end'))}")
    lines.append("")
    lines.append("## CONTENIDO")
    lines.append("A. ASPECTOS DE MAYOR INTERÉS")
    if isinstance(contenido.get("major_interest_titles"), list):
        for idx, title in enumerate(contenido.get("major_interest_titles"), start=1):
            lines.append(f"{idx}. {_safe_text(title)}")
    lines.append("B. ASPECTOS DE CONTROL INTERNO")
    if isinstance(contenido.get("control_titles"), list):
        for idx, title in enumerate(contenido.get("control_titles"), start=1):
            lines.append(f"{idx}. {_safe_text(title)}")
    lines.append("")
    lines.append(_safe_text(letter_intro.get("city_and_date")) or f"Quito, {_safe_text(header.get('date_label'))}")
    lines.append("")
    lines.append(_safe_text(letter_intro.get("recipient_block")) or f"A la Junta General de Accionistas,\n{_safe_text(header.get('company_name'))}")
    lines.append("")
    lines.append(_safe_text(letter_intro.get("saludo")) or "Estimados Señores:")
    lines.append("")
    paragraphs = letter_intro.get("paragraphs") if isinstance(letter_intro.get("paragraphs"), list) else []
    for p in paragraphs:
        lines.append(_safe_text(p))
        lines.append("")
    lines.append("## ASPECTOS DE MAYOR INTERÉS")
    if not major_findings:
        lines.append("[[PENDIENTE]] No se registraron hallazgos en esta sección.")
    for idx, finding in enumerate(major_findings, start=1):
        if not isinstance(finding, dict):
            continue
        lines.append("")
        lines.append(f"### {idx}. {_safe_text(finding.get('titulo'))}")
        antecedentes = _safe_text(finding.get("antecedentes"))
        if antecedentes:
            lines.append("Antecedentes:")
            lines.append(antecedentes)
        lines.append("Observación:")
        lines.append(_safe_text(finding.get("observacion")))
        lines.append("Recomendación:")
        lines.append(_safe_text(finding.get("recomendacion")))
        if include_response:
            lines.append("Comentarios de la Administración:")
            response = _safe_text(finding.get("comentarios_administracion")) or "[[PENDIENTE]]"
            lines.append(response)
    lines.append("")
    lines.append("## ASPECTOS DE CONTROL INTERNO")
    if not control_findings:
        lines.append("[[PENDIENTE]] No se registraron hallazgos en esta sección.")
    for idx, finding in enumerate(control_findings, start=1):
        if not isinstance(finding, dict):
            continue
        lines.append("")
        lines.append(f"### {idx}. {_safe_text(finding.get('titulo'))}")
        antecedentes = _safe_text(finding.get("antecedentes"))
        if antecedentes:
            lines.append("Antecedentes:")
            lines.append(antecedentes)
        lines.append("Observación:")
        lines.append(_safe_text(finding.get("observacion")))
        lines.append("Recomendación:")
        lines.append(_safe_text(finding.get("recomendacion")))
        if include_response:
            lines.append("Comentarios de la Administración:")
            response = _safe_text(finding.get("comentarios_administracion")) or "[[PENDIENTE]]"
            lines.append(response)
    lines.append("")
    lines.append(_safe_text(closing.get("body")) or _safe_text(document.get("closing")))
    lines.append("")
    lines.append(_safe_text(closing.get("signature_name")))
    lines.append(_safe_text(closing.get("signature_role")))
    lines.append(_safe_text(closing.get("firm_name")))
    lines.append(f"Cc: {_safe_text(closing.get('copy_to'))}")
    lines.append(_safe_text(closing.get("draft_legend")) or "BORRADOR")
    return "\n".join(lines).strip() + "\n"


def render_internal_control_letter_docx(document: dict[str, Any], *, contract: dict[str, Any]) -> tuple[bytes, str, str]:
    from docx.shared import Pt

    doc, template_version, template_mode = _load_docx_document("carta_control_interno.docx", contract=contract)
    header = document.get("header", {}) if isinstance(document.get("header"), dict) else {}
    cover = document.get("cover", {}) if isinstance(document.get("cover"), dict) else {}
    contenido = document.get("contenido", {}) if isinstance(document.get("contenido"), dict) else {}
    letter_intro = document.get("letter_intro", {}) if isinstance(document.get("letter_intro"), dict) else {}
    major_findings = document.get("major_interest_findings") if isinstance(document.get("major_interest_findings"), list) else []
    control_findings = document.get("internal_control_findings") if isinstance(document.get("internal_control_findings"), list) else []
    closing = document.get("closing", {}) if isinstance(document.get("closing"), dict) else {}
    include_response = bool(document.get("include_management_response"))

    title = doc.add_paragraph(_safe_text(cover.get("company_name")) or _safe_text(header.get("company_name")))
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    subtitle = doc.add_paragraph(_safe_text(cover.get("report_title")) or "Informe de Control Interno")
    subtitle.runs[0].bold = True
    doc.add_paragraph(f"Al {_safe_text(cover.get('period_end')) or _safe_text(header.get('period_end'))}")
    doc.add_page_break()

    doc.add_paragraph("CONTENIDO").runs[0].bold = True
    doc.add_paragraph("A. ASPECTOS DE MAYOR INTERÉS").runs[0].bold = True
    for idx, t in enumerate(contenido.get("major_interest_titles") if isinstance(contenido.get("major_interest_titles"), list) else [], start=1):
        doc.add_paragraph(f"{idx}. {_safe_text(t)}")
    doc.add_paragraph("B. ASPECTOS DE CONTROL INTERNO").runs[0].bold = True
    for idx, t in enumerate(contenido.get("control_titles") if isinstance(contenido.get("control_titles"), list) else [], start=1):
        doc.add_paragraph(f"{idx}. {_safe_text(t)}")
    doc.add_page_break()

    doc.add_paragraph(_safe_text(letter_intro.get("city_and_date")) or f"Quito, {_safe_text(header.get('date_label'))}")
    doc.add_paragraph(_safe_text(letter_intro.get("recipient_block")) or f"A la Junta General de Accionistas,\n{_safe_text(header.get('company_name'))}")
    doc.add_paragraph(_safe_text(letter_intro.get("saludo")) or "Estimados Señores:")
    for p_text in (letter_intro.get("paragraphs") if isinstance(letter_intro.get("paragraphs"), list) else []):
        doc.add_paragraph(_safe_text(p_text))
    doc.add_paragraph("")

    def _render_finding_block(idx: int, finding: dict[str, Any]) -> None:
        doc.add_paragraph(f"{idx}. {_safe_text(finding.get('titulo'))}").runs[0].bold = True
        antecedentes = _safe_text(finding.get("antecedentes"))
        if antecedentes:
            doc.add_paragraph("Antecedentes:").runs[0].bold = True
            doc.add_paragraph(antecedentes)
        doc.add_paragraph("Observación:").runs[0].bold = True
        doc.add_paragraph(_safe_text(finding.get("observacion")))
        doc.add_paragraph("Recomendación:").runs[0].bold = True
        doc.add_paragraph(_safe_text(finding.get("recomendacion")))
        if include_response:
            doc.add_paragraph("Comentarios de la Administración:").runs[0].bold = True
            doc.add_paragraph(_safe_text(finding.get("comentarios_administracion")) or "[[PENDIENTE]]")

    doc.add_paragraph("ASPECTOS DE MAYOR INTERÉS").runs[0].bold = True
    if not major_findings:
        doc.add_paragraph("[[PENDIENTE]] No se registraron hallazgos en esta sección.")
    for idx, finding in enumerate(major_findings, start=1):
        if isinstance(finding, dict):
            _render_finding_block(idx, finding)

    doc.add_paragraph("ASPECTOS DE CONTROL INTERNO").runs[0].bold = True
    if not control_findings:
        doc.add_paragraph("[[PENDIENTE]] No se registraron hallazgos en esta sección.")
    for idx, finding in enumerate(control_findings, start=1):
        if isinstance(finding, dict):
            _render_finding_block(idx, finding)

    doc.add_paragraph(_safe_text(closing.get("body")) or _safe_text(document.get("closing")))
    doc.add_paragraph(_safe_text(closing.get("signature_name")))
    doc.add_paragraph(_safe_text(closing.get("signature_role")))
    doc.add_paragraph(_safe_text(closing.get("firm_name")))
    doc.add_paragraph(f"Cc: {_safe_text(closing.get('copy_to'))}")
    draft = doc.add_paragraph(_safe_text(closing.get("draft_legend")) or "BORRADOR")
    if draft.runs:
        draft.runs[0].italic = True

    output = BytesIO()
    doc.save(output)
    return output.getvalue(), template_version, template_mode


def _build_niif_notes_payload(
    cliente_id: str,
    *,
    ifrs_for_smes_version: str,
    early_adoption: bool,
) -> dict[str, Any]:
    perfil = read_perfil(cliente_id) or {}
    cliente = perfil.get("cliente", {}) if isinstance(perfil.get("cliente"), dict) else {}
    encargo = perfil.get("encargo", {}) if isinstance(perfil.get("encargo"), dict) else {}
    operacion = perfil.get("operacion", {}) if isinstance(perfil.get("operacion"), dict) else {}
    contexto = perfil.get("contexto_negocio", {}) if isinstance(perfil.get("contexto_negocio"), dict) else {}
    materialidad = perfil.get("materialidad", {}) if isinstance(perfil.get("materialidad"), dict) else {}
    riesgo = perfil.get("riesgo_global", {}) if isinstance(perfil.get("riesgo_global"), dict) else {}
    notes_flags = {
        "partes_relacionadas": bool(contexto.get("tiene_partes_relacionadas")),
        "inventarios_significativos": bool(operacion.get("tiene_inventarios_significativos")),
        "cartera_significativa": bool(operacion.get("tiene_cartera_significativa")),
        "empleados": bool(operacion.get("tiene_empleados") or perfil.get("nomina", {}).get("tiene_empleados")),
    }
    return {
        "cliente": {
            "nombre_legal": _safe_text(cliente.get("nombre_legal")) or cliente_id,
            "nombre_corto": _safe_text(cliente.get("nombre_corto")),
            "pais": _safe_text(cliente.get("pais")) or "Ecuador",
            "domicilio": _safe_text(cliente.get("domicilio")) or "[[PENDIENTE]] domicilio legal",
            "sector": _safe_text(cliente.get("sector")) or "[[PENDIENTE]] sector",
            "moneda_funcional": _safe_text(cliente.get("moneda_funcional")) or "USD",
        },
        "encargo": {
            "anio_activo": encargo.get("anio_activo"),
            "periodo_inicio": _safe_text(encargo.get("periodo_inicio")) or "[[PENDIENTE]] periodo inicio",
            "periodo_fin": _safe_text(encargo.get("periodo_fin")) or "[[PENDIENTE]] periodo fin",
            "marco_referencial": _safe_text(encargo.get("marco_referencial")) or "NIIF_PYMES",
            "ifrs_for_smes_version": ifrs_for_smes_version,
            "early_adoption": bool(early_adoption),
        },
        "materialidad": materialidad,
        "riesgo_global": riesgo,
        "banderas_revelacion": notes_flags,
        "findings": collect_findings(cliente_id, max_findings=5),
    }


def build_niif_pymes_document(payload: dict[str, Any], *, include_policy_section: bool) -> dict[str, Any]:
    cliente = payload.get("cliente", {}) if isinstance(payload.get("cliente"), dict) else {}
    encargo = payload.get("encargo", {}) if isinstance(payload.get("encargo"), dict) else {}
    flags = payload.get("banderas_revelacion", {}) if isinstance(payload.get("banderas_revelacion"), dict) else {}
    materialidad = payload.get("materialidad", {}) if isinstance(payload.get("materialidad"), dict) else {}
    final_mat = materialidad.get("final", {}) if isinstance(materialidad.get("final"), dict) else {}
    prelim_mat = materialidad.get("preliminar", {}) if isinstance(materialidad.get("preliminar"), dict) else {}
    mat_ref = final_mat.get("materialidad_planeacion") or prelim_mat.get("materialidad_global") or "[[PENDIENTE]] materialidad"

    sections: list[dict[str, Any]] = [
        {
            "id": "informacion_general",
            "title": "Informacion general",
            "content": (
                f"La entidad {_safe_text(cliente.get('nombre_legal'))} se encuentra domiciliada en "
                f"{_safe_text(cliente.get('domicilio'))}, {_safe_text(cliente.get('pais'))}. "
                f"Su actividad principal corresponde al sector {_safe_text(cliente.get('sector'))}."
            ),
        },
        {
            "id": "base_preparacion",
            "title": "Base de preparacion",
            "content": (
                f"Los estados financieros para el periodo terminado al {_safe_text(encargo.get('periodo_fin'))} "
                f"se presentan como BORRADOR y se preparan con NIIF para PYMES (edicion "
                f"{_safe_text(encargo.get('ifrs_for_smes_version'))})"
                f"{' con adopcion anticipada' if bool(encargo.get('early_adoption')) else ''}. "
                f"No se afirma cumplimiento definitivo hasta validacion final del expediente."
            ),
        },
    ]
    if include_policy_section:
        sections.append(
            {
                "id": "politicas",
                "title": "Politicas contables significativas",
                "content": (
                    "Politicas en estado BORRADOR. "
                    "Si no existe manual formal, mantener [[PENDIENTE]] hasta aprobacion de administracion."
                ),
            }
        )
    sections.extend(
        [
            {
                "id": "juicios",
                "title": "Juicios y estimaciones significativas",
                "content": (
                    "La preparacion requiere juicios sobre deterioro, provisiones y revelaciones. "
                    "Cuando la evidencia sea incompleta, se marcara [[PENDIENTE]] en la nota respectiva."
                ),
            },
            {
                "id": "notas_especificas",
                "title": "Notas especificas sugeridas",
                "content": {
                    "partes_relacionadas": "Incluir" if bool(flags.get("partes_relacionadas")) else "No aplica",
                    "inventarios": "Incluir" if bool(flags.get("inventarios_significativos")) else "No aplica",
                    "cuentas_por_cobrar": "Incluir" if bool(flags.get("cartera_significativa")) else "Revisar materialidad",
                    "beneficios_empleados": "Incluir" if bool(flags.get("empleados")) else "No aplica",
                    "materialidad_referencia": mat_ref,
                },
            },
            {
                "id": "cierre",
                "title": "Cierre",
                "content": (
                    "Este documento es BORRADOR para revision interna. Debe validarse con TB, mayor, "
                    "soportes y revelaciones finales antes de emision."
                ),
            },
        ]
    )
    return {
        "header": {
            "title": "BORRADOR DE NOTAS A LOS ESTADOS FINANCIEROS - NIIF PARA PYMES",
            "company_name": _safe_text(cliente.get("nombre_legal")),
            "period_end": _safe_text(encargo.get("periodo_fin")),
            "ifrs_for_smes_version": _safe_text(encargo.get("ifrs_for_smes_version")),
            "early_adoption": bool(encargo.get("early_adoption")),
        },
        "sections": sections,
    }


def render_niif_pymes_markdown(document: dict[str, Any]) -> str:
    header = document.get("header", {}) if isinstance(document.get("header"), dict) else {}
    sections = document.get("sections", []) if isinstance(document.get("sections"), list) else []
    lines: list[str] = []
    lines.append(_safe_text(header.get("title")) or "BORRADOR DE NOTAS NIIF PYMES")
    lines.append("")
    lines.append(f"Entidad: {_safe_text(header.get('company_name'))}")
    lines.append(f"Periodo: {_safe_text(header.get('period_end'))}")
    lines.append(f"Version NIIF para PYMES: {_safe_text(header.get('ifrs_for_smes_version'))}")
    lines.append(f"Adopcion anticipada: {'si' if bool(header.get('early_adoption')) else 'no'}")
    lines.append("")
    for sec in sections:
        if not isinstance(sec, dict):
            continue
        lines.append(f"## {_safe_text(sec.get('title'))}")
        content = sec.get("content")
        if isinstance(content, dict):
            for key, value in content.items():
                lines.append(f"- {key}: {_safe_text(value)}")
        else:
            lines.append(_safe_text(content))
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def render_niif_pymes_docx(document: dict[str, Any], *, contract: dict[str, Any]) -> tuple[bytes, str, str]:
    from docx.shared import RGBColor

    doc, template_version, template_mode = _load_docx_document("notas_niif_pymes.docx", contract=contract)
    header = document.get("header", {}) if isinstance(document.get("header"), dict) else {}
    sections = document.get("sections", []) if isinstance(document.get("sections"), list) else []

    title = doc.add_paragraph()
    run = title.add_run(_safe_text(header.get("title")) or "BORRADOR DE NOTAS NIIF PYMES")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
    doc.add_paragraph(f"Entidad: {_safe_text(header.get('company_name'))}")
    doc.add_paragraph(f"Periodo: {_safe_text(header.get('period_end'))}")
    doc.add_paragraph(f"Version NIIF para PYMES: {_safe_text(header.get('ifrs_for_smes_version'))}")
    doc.add_paragraph(f"Adopcion anticipada: {'si' if bool(header.get('early_adoption')) else 'no'}")

    doc.add_page_break()
    doc.add_heading("Indice (borrador)", level=1)
    for sec in sections:
        if not isinstance(sec, dict):
            continue
        doc.add_paragraph(f"- {_safe_text(sec.get('title'))}")
    doc.add_page_break()

    for sec in sections:
        if not isinstance(sec, dict):
            continue
        doc.add_heading(_safe_text(sec.get("title")), level=2)
        content = sec.get("content")
        if isinstance(content, dict):
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Concepto"
            table.rows[0].cells[1].text = "Detalle"
            for key, value in content.items():
                row = table.add_row().cells
                row[0].text = _safe_text(key)
                row[1].text = _safe_text(value)
        else:
            doc.add_paragraph(_safe_text(content))

    footer = doc.add_paragraph("Estado del documento: BORRADOR")
    footer.runs[0].italic = True

    output = BytesIO()
    doc.save(output)
    return output.getvalue(), template_version, template_mode


def generate_internal_control_letter(
    cliente_id: str,
    *,
    recipient: str = "Gerencia General",
    include_management_response: bool = True,
    max_findings: int = 10,
    requested_by: str = "",
) -> dict[str, Any]:
    document_type = "carta_control_interno"
    contract = _load_template_contract(document_type, cliente_id=cliente_id)
    generated_at = datetime.now(timezone.utc)
    normalized_recipient = _normalize_recipient(recipient)
    perfil = read_perfil(cliente_id) or {}
    cliente = perfil.get("cliente", {}) if isinstance(perfil.get("cliente"), dict) else {}
    encargo = perfil.get("encargo", {}) if isinstance(perfil.get("encargo"), dict) else {}
    company_name = _safe_text(cliente.get("nombre_legal")) or cliente_id
    period_end = _safe_text(encargo.get("periodo_fin")) or f"{encargo.get('anio_activo', '')}-12-31"
    ruc = _safe_text(cliente.get("ruc")) or _safe_text(cliente.get("identificacion")) or "[[PENDIENTE]]"
    findings = collect_findings(cliente_id, max_findings=max_findings)

    document = build_internal_control_letter(
        company_name=company_name,
        period_end=period_end,
        recipient=normalized_recipient,
        findings=findings,
        include_management_response=include_management_response,
    )
    cover = document.get("cover") if isinstance(document.get("cover"), dict) else {}
    if isinstance(cover, dict):
        cover["ruc"] = ruc
    content = render_internal_control_letter_markdown(document)
    input_payload = {
        "recipient": normalized_recipient,
        "include_management_response": include_management_response,
        "max_findings": max_findings,
        "findings_count": len(findings),
    }
    generation_metadata = _generation_meta_fallback(
        generated_at=generated_at,
        requested_by=requested_by,
        input_payload=input_payload,
        document_type=document_type,
        template_mode="default",
        template_version=_safe_text(contract.get("template_version")) or "v1-default",
        contract=contract,
    )
    generation_metadata["template_locked"] = True

    query = (
        "Redacta carta de control interno en espanol corporativo de Ecuador.\n"
        "No inventes normas, cifras, ni afirmes cumplimiento definitivo.\n"
        f"Data estructurada: {document}"
    )
    try:
        rag = generate_judgement_response(cliente_id, query, mode="carta_control_interno")
        llm_text = _safe_text(rag.get("answer"))
        if len(llm_text) >= 350:
            content = llm_text + ("\n" if not llm_text.endswith("\n") else "")
            generation_metadata = _generation_meta_llm(
                rag=rag,
                generated_at=generated_at,
                requested_by=requested_by,
                input_payload=input_payload,
                document_type=document_type,
                template_mode=_safe_text(generation_metadata.get("template_mode")) or "default",
                template_version=_safe_text(generation_metadata.get("template_version")) or "v1-default",
                contract=contract,
            )
            generation_metadata["template_locked"] = True
    except Exception:
        pass

    ts = generated_at.strftime("%Y%m%dT%H%M%SZ")
    out_dir = _report_dir(cliente_id)
    md_path = out_dir / f"carta_control_interno_{ts}.md"
    md_bytes = content.encode("utf-8")
    md_path.write_bytes(md_bytes)
    try:
        docx_bytes, template_version, template_mode = render_internal_control_letter_docx(document, contract=contract)
    except Exception:
        docx_bytes = md_bytes
        template_version = "v1-fallback-no-docx"
        template_mode = "fallback"
    docx_path = out_dir / f"carta_control_interno_{ts}.docx"
    docx_path.write_bytes(docx_bytes)

    artifacts = [
        _build_artifact(
            artifact_type="markdown",
            artifact_path=md_path,
            payload=md_bytes,
            template_version=template_version,
        ),
        _build_artifact(
            artifact_type="docx",
            artifact_path=docx_path,
            payload=docx_bytes,
            template_version=template_version,
        ),
    ]
    generation_metadata["template_mode"] = template_mode
    generation_metadata["template_version"] = template_version
    generation_metadata["template_locked"] = True
    return {
        "content": content,
        "document": document,
        "path": str(docx_path),
        "findings_count": len(findings),
        "recipient": normalized_recipient,
        "generation_metadata": generation_metadata,
        "artifacts": artifacts,
    }


def generate_niif_pymes_draft(
    cliente_id: str,
    *,
    ifrs_for_smes_version: str = "2025",
    early_adoption: bool = False,
    include_policy_section: bool = True,
    requested_by: str = "",
) -> dict[str, Any]:
    document_type = "niif_pymes_borrador"
    contract = _load_template_contract(document_type, cliente_id=cliente_id)
    generated_at = datetime.now(timezone.utc)
    payload = _build_niif_notes_payload(
        cliente_id,
        ifrs_for_smes_version=ifrs_for_smes_version,
        early_adoption=early_adoption,
    )
    document = build_niif_pymes_document(payload, include_policy_section=include_policy_section)
    content = render_niif_pymes_markdown(document)
    input_payload = {
        "ifrs_for_smes_version": ifrs_for_smes_version,
        "early_adoption": early_adoption,
        "include_policy_section": include_policy_section,
    }
    generation_metadata = _generation_meta_fallback(
        generated_at=generated_at,
        requested_by=requested_by,
        input_payload=input_payload,
        document_type=document_type,
        template_mode="default",
        template_version=_safe_text(contract.get("template_version")) or "v1-default",
        contract=contract,
    )

    query = (
        "Genera borrador de notas NIIF para PYMES.\n"
        "No inventes cifras ni hechos. Si falta data, conserva [[PENDIENTE]].\n"
        f"Data estructurada: {document}"
    )
    try:
        rag = generate_judgement_response(cliente_id, query, mode="niif_notas")
        llm_text = _safe_text(rag.get("answer"))
        if len(llm_text) >= 450:
            content = llm_text + ("\n" if not llm_text.endswith("\n") else "")
            generation_metadata = _generation_meta_llm(
                rag=rag,
                generated_at=generated_at,
                requested_by=requested_by,
                input_payload=input_payload,
                document_type=document_type,
                template_mode=_safe_text(generation_metadata.get("template_mode")) or "default",
                template_version=_safe_text(generation_metadata.get("template_version")) or "v1-default",
                contract=contract,
            )
    except Exception:
        pass

    ts = generated_at.strftime("%Y%m%dT%H%M%SZ")
    period = _safe_text((payload.get("encargo") or {}).get("periodo_fin")) or "periodo"
    safe_period = period.replace("-", "")
    out_dir = _report_dir(cliente_id)
    md_path = out_dir / f"notas_niif_pymes_{safe_period}_{ts}.md"
    md_bytes = content.encode("utf-8")
    md_path.write_bytes(md_bytes)
    try:
        docx_bytes, template_version, template_mode = render_niif_pymes_docx(document, contract=contract)
    except Exception:
        docx_bytes = md_bytes
        template_version = "v1-fallback-no-docx"
        template_mode = "fallback"
    docx_path = out_dir / f"notas_niif_pymes_{safe_period}_{ts}.docx"
    docx_path.write_bytes(docx_bytes)

    artifacts = [
        _build_artifact(
            artifact_type="markdown",
            artifact_path=md_path,
            payload=md_bytes,
            template_version=template_version,
        ),
        _build_artifact(
            artifact_type="docx",
            artifact_path=docx_path,
            payload=docx_bytes,
            template_version=template_version,
        ),
    ]
    generation_metadata["template_mode"] = template_mode
    generation_metadata["template_version"] = template_version
    return {
        "content": content,
        "document": document,
        "path": str(docx_path),
        "period_end": period,
        "ifrs_for_smes_version": str(ifrs_for_smes_version),
        "early_adoption": bool(early_adoption),
        "generation_metadata": generation_metadata,
        "artifacts": artifacts,
    }
